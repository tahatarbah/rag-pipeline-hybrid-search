from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from app.config import settings

SUPPORTED_SUFFIXES = {".md", ".txt", ".pdf", ".docx"}
SAFE_NAME = re.compile(r"[^A-Za-z0-9._-]+")


def safe_filename(name: str) -> str:
    base = Path(name).name
    cleaned = SAFE_NAME.sub("_", base).strip("._")
    return cleaned or "upload.bin"


@dataclass
class Document:
    path: Path
    source: str
    title: str
    text: str
    page: int | None = None
    extra: dict = field(default_factory=dict)


def _title_from_text(text: str, fallback: str) -> str:
    for line in text.splitlines():
        stripped = line.strip().lstrip("#").strip()
        if stripped:
            return stripped[:200]
    return fallback


def _read_text_file(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8", errors="replace")
    source = path.name
    return [
        Document(
            path=path,
            source=source,
            title=_title_from_text(text, source),
            text=text,
            page=None,
        )
    ]


def _read_pdf(path: Path) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    source = path.name
    docs: list[Document] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        docs.append(
            Document(
                path=path,
                source=source,
                title=_title_from_text(text, f"{source} p.{i}"),
                text=text,
                page=i,
            )
        )
    return docs


def _read_docx(path: Path) -> list[Document]:
    from docx import Document as DocxDocument

    parsed = DocxDocument(str(path))
    parts = [p.text for p in parsed.paragraphs if p.text and p.text.strip()]
    text = "\n\n".join(parts)
    source = path.name
    return [
        Document(
            path=path,
            source=source,
            title=_title_from_text(text, source),
            text=text,
            page=None,
        )
    ]


def load_file(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return _read_text_file(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    return []


def load_documents(docs_dir: Path | None = None) -> list[Document]:
    root = docs_dir or settings.docs_dir
    if not root.exists():
        return []

    documents: list[Document] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if path.name.startswith("."):
            continue
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        documents.extend(load_file(path))
    return documents
